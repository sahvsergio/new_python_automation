import docx
from docx.shared import Inches
#opening the file
doc = docx.Document('SergioAndrésHerrera Velásquez_Anexo_guia_aap1.docx')
print('Document Object:', doc)


#getting info out of the doc
print('Title of the document')
print(doc.paragraphs[1].text)

#attributes of the doc

print('Attributes of the document')
print('Author:',doc.core_properties.author)
print('Date Created:',doc.core_properties.created)
print('Document Revision',doc.core_properties.revision)

#reading tables

table=doc.tables[0]
print(table)
print('Column 1')
for i in range(len(table.rows)):
    print(table.rows[i].cells[0].paragraphs[0].text)
    

print('Column 2')
for i in range(len(table.rows)):
    print(table.rows[i].cells[1].paragraphs[0].text)
    
#writing data into  word Document(Adding headings, images,tables)
from docx import Document
#document object
document=Document()
#adding headings
document.add_heading('Test Document from Docs',0)
document.add_heading('Lets talk about python language', level=1 )
#adding paragraphs
p= document.add_paragraph('A plai paragraph having some')
p.add_run('bold words').bold=True#adding style to paragraphs
p.add_run('and  italics').italic=True
document.add_paragraph('First lets see the Python logo', style='List Bullet')#adding bulleted lists
#adding pictures
document.add_picture('favicon.png', width=Inches(6.25))
#adding pictures
table=document.add_table(rows=1,cols=3)
table.style='Table Grid'
data={'id':1,'items':'apple','price':50}
headings=table.rows[0].cells
headings[0].text='Id'
headings[1].text='items'
headings[2].text= 'Price'

row=table.add_row().cells
row[0].text=str(data.get('id'))
row[1].text=data.get('items')
row[2].text=str(data.get('price'))


#saving te document
document.save('testdoc.docx')
    
